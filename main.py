#!/usr/bin/python
#SixH
import sqlite3
from docx import Document
from docx.shared import Inches

from docx.oxml.ns import qn

department =''
name = ''
pid =''
date = ''
diary = ''
document = Document()
conn = sqlite3.connect('d:/Workspace/Python/Diary.db')
c = conn.cursor()
print("Opened database successfully")
cursor = c.execute("SELECT * FROM Diary")
for row in cursor:
    department = "轮岗部门: " + row[1]
    name = "员工名字: " + row[2]
    pid = "员工工号: " + row[3]
    date = "周记日期: " + row[4]
    diary = "周记内容: \n" + row[5]
    print("轮岗部门 = ", department)
    print("员工名字 = ", name)
    print("员工工号 = ", pid)
    print("日期 = ", date)
    print("周记内容 = ", diary, "\n")
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.add_heading(department, 0)
    document.add_paragraph(name, style='List Bullet')
    document.add_paragraph(pid, style='List Bullet')
    document.add_paragraph(date, style='List Bullet')
    document.add_paragraph(diary, style='List Bullet')
    document.add_page_break()

conn.close()
document.save('demo.docx')
